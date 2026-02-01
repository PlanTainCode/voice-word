import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def generate_word_document(title: str, text: str, output_dir: str) -> str:
    """Генерация Word документа с отформатированным текстом"""
    
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # Заголовок
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Дата создания
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    date_run.font.size = Pt(10)
    date_run.font.italic = True
    
    doc.add_paragraph()  # Пустая строка
    
    # Основной текст - разбиваем по абзацам
    paragraphs = text.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Первая строка с отступом
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.paragraph_format.space_after = Pt(12)
            
            run = para.add_run(para_text.strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
    
    # Сохраняем файл
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
    filename = f"{safe_title}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)
    
    os.makedirs(output_dir, exist_ok=True)
    doc.save(filepath)
    
    return filepath

